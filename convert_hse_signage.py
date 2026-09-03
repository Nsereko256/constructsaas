from pathlib import Path
import subprocess
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from pypdf import PdfReader


SOURCE = Path(r"C:\Users\840G3\Downloads\HSE SIGNAGE.pdf")
OUTPUT = Path(r"C:\Users\840G3\Desktop\CONSTRUCT SAAS\HSE SIGNAGE.docx")
WORK = Path(r"C:\Users\840G3\Desktop\CONSTRUCT SAAS\tmp_hse_signage")
PDFTOPPM = Path(r"C:\Users\840G3\.cache\codex-runtimes\codex-primary-runtime\dependencies\native\poppler\Library\bin\pdftoppm.exe")


def configure_section(section, width_pt, height_pt):
    section.page_width = Inches(width_pt / 72)
    section.page_height = Inches(height_pt / 72)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)


def main():
    reader = PdfReader(str(SOURCE))
    WORK.mkdir(parents=True, exist_ok=True)
    prefix = WORK / "page"
    subprocess.run(
        [str(PDFTOPPM), "-jpeg", "-r", "150", "-jpegopt", "quality=95", str(SOURCE), str(prefix)],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )

    image_paths = sorted(WORK.glob("page-*.jpg"), key=lambda p: int(p.stem.split("-")[-1]))
    if len(image_paths) != len(reader.pages):
        raise RuntimeError(f"Expected {len(reader.pages)} rendered pages, found {len(image_paths)}")

    doc = Document()
    first_page = reader.pages[0]
    configure_section(doc.sections[0], float(first_page.mediabox.width), float(first_page.mediabox.height))
    doc.add_paragraph()

    for index, (page, image_path) in enumerate(zip(reader.pages, image_paths)):
        if index:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            configure_section(section, float(page.mediabox.width), float(page.mediabox.height))
            paragraph = doc.add_paragraph()
        else:
            paragraph = doc.paragraphs[-1]
        paragraph.paragraph_format.space_before = Inches(0)
        paragraph.paragraph_format.space_after = Inches(0)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.add_run().add_picture(
            str(image_path),
            width=Inches(float(page.mediabox.width) / 72),
            height=Inches(float(page.mediabox.height) / 72),
        )

    doc.core_properties.title = "HSE SIGNAGE"
    doc.core_properties.subject = "Converted from PDF with one image per page"
    doc.save(OUTPUT)

    shutil.rmtree(WORK, ignore_errors=True)
    print(f"Created {OUTPUT}")
    print(f"Pages: {len(reader.pages)}")


if __name__ == "__main__":
    main()
