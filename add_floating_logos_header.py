from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT = Path(r"C:\Users\840G3\Downloads\HSE SIGNAGE (1).docx")
OUTPUT = Path(r"C:\Users\840G3\Desktop\CONSTRUCT SAAS\HSE SIGNAGE - Egypro ATC headers.docx")
EGYPRO = Path(r"C:\Users\840G3\Downloads\WhatsApp Image 2026-08-25 at 11.14.06.jpeg")
ATC = Path(r"C:\Users\840G3\Downloads\WhatsApp Image 2026-08-25 at 11.13.37.jpeg")


def emu(inches):
    return str(int(round(inches * 914400)))


def floating_picture(paragraph, image_path, width_inches, x_inches):
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    inline = run._r.xpath('.//wp:inline')[0]
    extent = inline.find(qn('wp:extent'))
    doc_pr = inline.find(qn('wp:docPr'))
    c_nv = inline.find(qn('wp:cNvGraphicFramePr'))
    graphic = inline.find(qn('a:graphic'))

    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', '251659264')
    anchor.set('behindDoc', '0')
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')

    simple = OxmlElement('wp:simplePos')
    simple.set('x', '0')
    simple.set('y', '0')
    anchor.append(simple)

    pos_h = OxmlElement('wp:positionH')
    pos_h.set('relativeFrom', 'page')
    off_h = OxmlElement('wp:posOffset')
    off_h.text = emu(x_inches)
    pos_h.append(off_h)
    anchor.append(pos_h)

    pos_v = OxmlElement('wp:positionV')
    pos_v.set('relativeFrom', 'page')
    off_v = OxmlElement('wp:posOffset')
    off_v.text = emu(0.12)
    pos_v.append(off_v)
    anchor.append(pos_v)

    anchor.append(deepcopy(extent))
    effect = OxmlElement('wp:effectExtent')
    for side in ('l', 't', 'r', 'b'):
        effect.set(side, '0')
    anchor.append(effect)
    anchor.append(OxmlElement('wp:wrapNone'))
    anchor.append(deepcopy(doc_pr))
    anchor.append(deepcopy(c_nv))
    anchor.append(deepcopy(graphic))
    inline.getparent().replace(inline, anchor)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn('w:pPr'):
            paragraph._p.remove(child)


doc = Document(str(INPUT))
header = doc.sections[0].header
clear_paragraph(header.paragraphs[0])
p = header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Inches(0)
p.paragraph_format.space_after = Inches(0)
floating_picture(p, EGYPRO, 0.80, 0.25)
floating_picture(p, ATC, 1.20, 15.05)

doc.save(str(OUTPUT))
print(OUTPUT)
