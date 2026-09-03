from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

INPUT = Path(r"C:\Users\840G3\Downloads\HSE SIGNAGE (1).docx")
OUTPUT = Path(r"C:\Users\840G3\Desktop\CONSTRUCT SAAS\HSE SIGNAGE - Egypro ATC headers.docx")
EGYPRO = Path(r"C:\Users\840G3\Downloads\WhatsApp Image 2026-08-25 at 11.14.06.jpeg")
ATC = Path(r"C:\Users\840G3\Downloads\WhatsApp Image 2026-08-25 at 11.13.37.jpeg")


def remove_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def configure_header(header):
    # Keep the existing header structure minimal while preserving section linkage.
    while len(header.paragraphs) > 1:
        p = header.paragraphs[-1]._element
        p.getparent().remove(p)
    first = header.paragraphs[0]
    clear_paragraph(first)
    first.paragraph_format.space_before = Inches(0)
    first.paragraph_format.space_after = Inches(0)

    table = header.add_table(rows=1, cols=2, width=Inches(16.05))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_borders(table)
    table.rows[0].height = Inches(0.72)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    left, right = table.rows[0].cells
    left.width = Inches(8.025)
    right.width = Inches(8.025)
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for side in ("top", "left", "bottom", "right"):
            node = tc_mar.find(qn("w:" + side))
            if node is None:
                node = OxmlElement("w:" + side)
                tc_mar.append(node)
            node.set(qn("w:w"), "0")
            node.set(qn("w:type"), "dxa")

    lp = left.paragraphs[0]
    rp = right.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    lp.paragraph_format.space_before = Inches(0)
    lp.paragraph_format.space_after = Inches(0)
    rp.paragraph_format.space_before = Inches(0)
    rp.paragraph_format.space_after = Inches(0)
    lp.add_run().add_picture(str(EGYPRO), width=Inches(0.62))
    rp.add_run().add_picture(str(ATC), width=Inches(0.88))


doc = Document(str(INPUT))
for section in doc.sections:
    # Keep the header inside the existing top margin so the signage artwork
    # retains the original pagination and placement.
    section.header_distance = Inches(0.08)
configure_header(doc.sections[0].header)

# Usually all sections inherit the first header. Ensure any independent or
# special first/even-page headers also receive the same branding.
for index, section in enumerate(doc.sections):
    if index > 0 and not section.header.is_linked_to_previous:
        configure_header(section.header)
    if section.different_first_page_header_footer:
        configure_header(section.first_page_header)
        configure_header(section.even_page_header)

doc.save(str(OUTPUT))
print(OUTPUT)
